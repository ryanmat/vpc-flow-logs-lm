# Description: Generates customer-facing .docx deployment guide for GCP VPC Flow Logs.
# Description: Run with: uv run python generate_customer_docs.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_step(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    return table


def build_deployment_guide():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    doc.add_heading("GCP VPC Flow Logs to LogicMonitor", level=0)
    doc.add_paragraph("Deployment and Configuration Guide")
    doc.add_paragraph("")

    # --- Prerequisites ---
    add_heading(doc, "1. Prerequisites", 1)

    add_heading(doc, "GCP", 2)
    add_bullet(doc, "A GCP project with billing enabled")
    add_bullet(doc, "A Pub/Sub topic for VPC flow logs (already created)")
    add_bullet(doc, "VPC subnets with Flow Logs enabled")
    add_bullet(doc, "gcloud CLI installed and authenticated")
    add_bullet(doc, "IAM roles: Cloud Functions Admin, Pub/Sub Admin, Logging Admin, Secret Manager Admin, Service Account User")

    add_heading(doc, "LogicMonitor", 2)
    add_bullet(doc, "LM portal with LM Logs enabled")
    add_bullet(doc, "GCP project added as a Cloud Account in LM (for device auto-discovery)")
    add_bullet(doc, "An API Only User with Manage permission for Logs & Traces")
    add_bullet(doc, "A Bearer Token generated for that user")

    # --- Log Router Sink ---
    add_heading(doc, "2. Create Log Router Sink", 1)
    doc.add_paragraph(
        "The Log Router sink exports VPC Flow Logs from Cloud Logging to your Pub/Sub topic."
    )

    doc.add_paragraph(
        'gcloud logging sinks create vpc-flowlogs-to-lm \\\n'
        '    "pubsub.googleapis.com/projects/PROJECT_ID/topics/TOPIC_NAME" \\\n'
        '    --project=PROJECT_ID \\\n'
        '    --log-filter=\'resource.type="gce_subnetwork" '
        'log_id("compute.googleapis.com/vpc_flows")\' \\\n'
        '    --quiet',
        style="No Spacing",
    )
    doc.add_paragraph("")
    doc.add_paragraph("Grant the sink's service account Pub/Sub Publisher on the topic:")
    doc.add_paragraph(
        "gcloud pubsub topics add-iam-policy-binding TOPIC_NAME \\\n"
        '    --member="SINK_WRITER_IDENTITY" \\\n'
        '    --role="roles/pubsub.publisher"',
        style="No Spacing",
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Get the writer identity with: "
        'gcloud logging sinks describe vpc-flowlogs-to-lm --format="value(writerIdentity)"'
    )

    # --- Secret Manager ---
    add_heading(doc, "3. Create Secret Manager Entries", 1)
    doc.add_paragraph("Store LM credentials in GCP Secret Manager:")

    add_table(doc, ["Secret Name", "Value"], [
        ["lm-company-name", "Your LM portal subdomain (e.g., acmecorp)"],
        ["lm-bearer-token", "Bearer token from LM API Only User"],
    ])

    doc.add_paragraph("")
    doc.add_paragraph(
        'echo -n "YOUR_PORTAL_NAME" | gcloud secrets create lm-company-name \\\n'
        "    --replication-policy=automatic --data-file=-",
        style="No Spacing",
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        'echo -n "YOUR_BEARER_TOKEN" | gcloud secrets create lm-bearer-token \\\n'
        "    --replication-policy=automatic --data-file=-",
        style="No Spacing",
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Grant the default compute service account Secret Manager Accessor role:"
    )
    doc.add_paragraph(
        "gcloud projects add-iam-policy-binding PROJECT_ID \\\n"
        '    --member="serviceAccount:PROJECT_NUMBER-compute@developer.gserviceaccount.com" \\\n'
        '    --role="roles/secretmanager.secretAccessor"',
        style="No Spacing",
    )

    # --- Deploy Cloud Function ---
    add_heading(doc, "4. Deploy Cloud Function", 1)
    doc.add_paragraph("Deploy from the cloud_function/ directory:")
    doc.add_paragraph(
        "gcloud functions deploy vpc-flowlogs-to-lm \\\n"
        "    --gen2 \\\n"
        "    --project=PROJECT_ID \\\n"
        "    --region=REGION \\\n"
        "    --source=./cloud_function \\\n"
        "    --runtime=python312 \\\n"
        "    --entry-point=handle_pubsub \\\n"
        "    --trigger-topic=TOPIC_NAME \\\n"
        "    --memory=256Mi \\\n"
        "    --timeout=60s \\\n"
        "    --concurrency=1 \\\n"
        "    --min-instances=0 \\\n"
        "    --max-instances=100 \\\n"
        '    --set-env-vars="USE_WEBHOOK=true,WEBHOOK_SOURCE_NAME=GCP-VPC-FlowLogs" \\\n'
        '    --set-secrets="LM_COMPANY_NAME=lm-company-name:latest,'
        'LM_BEARER_TOKEN=lm-bearer-token:latest"',
        style="No Spacing",
    )
    doc.add_paragraph("")
    doc.add_paragraph("Verify deployment:")
    doc.add_paragraph(
        "gcloud functions describe vpc-flowlogs-to-lm \\\n"
        "    --project=PROJECT_ID --region=REGION --gen2",
        style="No Spacing",
    )

    # --- Configure Webhook LogSource ---
    add_heading(doc, "5. Configure Webhook LogSource in LM Portal", 1)

    add_heading(doc, "5.1 Create the LogSource", 2)
    add_step(doc, "Navigate to Modules > My Module Toolbox")
    add_step(doc, "Click Add > LogSource")
    add_step(doc, 'Select "Webhook" as the type')
    add_step(doc, 'Set Name to "GCP VPC Flow Logs", Group to "GCP"')

    add_heading(doc, "5.2 Filter (Match All Conditions)", 2)
    add_table(doc, ["Attribute", "Operation", "Value"], [
        ["SourceName", "Equal", "GCP-VPC-FlowLogs"],
    ])

    add_heading(doc, "5.3 Log Fields (all Webhook Attribute method)", 2)
    doc.add_paragraph(
        "Important: LM webhook LogSources only process string values as metadata. "
        "The Cloud Function stringifies all payload values for compatibility."
    )
    add_table(doc, ["Key", "Method", "Value"], [
        ["src_ip", "Webhook Attribute", "src_ip"],
        ["dest_ip", "Webhook Attribute", "dest_ip"],
        ["src_port", "Webhook Attribute", "src_port"],
        ["dest_port", "Webhook Attribute", "dest_port"],
        ["protocol", "Webhook Attribute", "protocol"],
        ["bytes_sent", "Webhook Attribute", "bytes_sent"],
        ["packets_sent", "Webhook Attribute", "packets_sent"],
        ["reporter", "Webhook Attribute", "reporter"],
        ["vm_name", "Webhook Attribute", "vm_name"],
        ["log_level", "Webhook Attribute", "Level"],
        ["resource_type", "Webhook Attribute", "resourceType"],
    ])

    add_heading(doc, "5.4 Resource Mapping (Match Any / OR)", 2)
    add_table(doc, ["Method", "Key", "Value"], [
        ["Webhook Attribute", "system.gcp.resourcename", "vm_name"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Important: Use system.gcp.resourcename (not system.hostname). "
        "GCP cloud-discovered devices have a composite system.hostname that will not "
        "match the simple VM name from flow logs. The GCP project must be added as a "
        "Cloud Account in LM for device auto-discovery."
    )

    add_step(doc, "Click Save")

    # --- Verify ---
    add_heading(doc, "6. Verify End-to-End", 1)
    add_step(doc, 'Navigate to Logs in the LM portal')
    add_step(doc, 'Search for: sourceName="GCP-VPC-FlowLogs"')
    add_step(doc, "Confirm logs are appearing with timestamps within the last few minutes")
    add_step(doc, "Click a log entry and verify:")
    add_bullet(doc, "Resource column shows the GCP VM name")
    add_bullet(doc, "Resource Type shows GCP/ComputeEngine")
    add_bullet(doc, "Log Level shows info")
    add_bullet(doc, "Tags are populated (src_ip, dest_ip, protocol, vm_name, etc.)")

    add_heading(doc, "Check Cloud Function logs:", 2)
    doc.add_paragraph(
        "gcloud functions logs read vpc-flowlogs-to-lm \\\n"
        "    --project=PROJECT_ID --region=REGION --gen2 --limit=20",
        style="No Spacing",
    )
    doc.add_paragraph("")
    doc.add_paragraph('Look for: LM_WEBHOOK status=202 body=Accepted')

    # --- Troubleshooting ---
    add_heading(doc, "7. Troubleshooting", 1)
    add_table(doc, ["Symptom", "Cause", "Fix"], [
        [
            "No logs in LM portal",
            "Cloud Function not receiving messages",
            "Check Cloud Function logs. Verify Log Router sink and Pub/Sub topic.",
        ],
        [
            "Logs appear under default logsource",
            "Non-string payload values or LogSource filter mismatch",
            "Verify all payload values are strings. Verify SourceName Equal filter.",
        ],
        [
            "Logs not mapped to devices",
            "Resource mapping key mismatch",
            "Use system.gcp.resourcename (not system.hostname). Ensure GCP Cloud Account is added to LM.",
        ],
        [
            "HTTP 202 but no logs visible",
            "LogSource resource mapping failing silently",
            "Check that vm_name is present in payload and matches system.gcp.resourcename on the device.",
        ],
        [
            "Function errors on deploy",
            "Secret Manager permissions missing",
            "Grant secretmanager.secretAccessor to the compute service account.",
        ],
    ])

    return doc


if __name__ == "__main__":
    doc = build_deployment_guide()
    out_path = "/home/rmatuszewski/dev/tools/vpc-flow-logs-lm/gcp/vpc-flow-logs/documentation/GCP_VPC_Flow_Logs_LM_Deployment_Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
